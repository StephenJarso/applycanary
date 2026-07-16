"""Resume extraction.

Pulls out both text *and* layout signals, because real ATS parsers fail on
structure far more often than on wording: a two-column layout gets read in the
wrong order, table cells get flattened or dropped, and text in headers/footers is
frequently ignored entirely. The rule engine in `app/pipeline/ats_rules.py`
consumes these signals.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

log = logging.getLogger(__name__)

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}

# Headings ATS parsers reliably recognise. Creative substitutes ("Where I've
# Made An Impact") often cause the whole block to be misfiled.
CANONICAL_SECTIONS = {
    "experience": (
        "experience", "work experience", "professional experience",
        "employment", "employment history", "work history", "career history",
    ),
    "education": ("education", "academic background", "qualifications"),
    "skills": ("skills", "technical skills", "core competencies", "competencies"),
    "projects": ("projects", "personal projects", "selected projects"),
    "summary": ("summary", "professional summary", "profile", "objective", "about"),
    "certifications": ("certifications", "certificates", "licenses"),
}

_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?){2,4}\d{2,4}")
_URL = re.compile(r"https?://[^\s)>\]]+|(?:www\.|linkedin\.com/|github\.com/)[^\s)>\]]+")


@dataclass(slots=True)
class ParsedResume:
    path: str = ""
    file_type: str = ""
    text: str = ""
    pages: int = 0

    # --- layout signals ---
    has_text_layer: bool = True          # False => scanned image, unreadable by ATS
    table_count: int = 0
    has_multi_column: bool = False
    header_footer_text: list[str] = field(default_factory=list)
    image_count: int = 0
    fonts: set[str] = field(default_factory=set)
    bullet_glyphs: set[str] = field(default_factory=set)

    # --- content signals ---
    sections_found: dict[str, str] = field(default_factory=dict)
    unknown_headings: list[str] = field(default_factory=list)
    emails: list[str] = field(default_factory=list)
    phones: list[str] = field(default_factory=list)
    urls: list[str] = field(default_factory=list)
    parse_errors: list[str] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return len(self.text.split())

    @property
    def words(self) -> list[str]:
        return re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-]*", self.text.lower())


def parse_resume(path: str | Path) -> ParsedResume:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"resume not found: {p}")
    ext = p.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"unsupported resume type {ext!r}; use one of {sorted(SUPPORTED)}")

    if ext == ".pdf":
        parsed = _parse_pdf(p)
    elif ext == ".docx":
        parsed = _parse_docx(p)
    else:
        parsed = _parse_text(p)

    _extract_content_signals(parsed)
    return parsed


# ---------------------------------------------------------------- pdf


def _parse_pdf(p: Path) -> ParsedResume:
    import pdfplumber

    out = ParsedResume(path=str(p), file_type="pdf")
    chunks: list[str] = []

    with pdfplumber.open(str(p)) as pdf:
        out.pages = len(pdf.pages)
        for page in pdf.pages:
            try:
                chunks.append(page.extract_text() or "")
                words = page.extract_words() or []
                try:
                    out.table_count += len(page.find_tables() or [])
                except Exception:  # noqa: BLE001 - table finder is best-effort
                    pass
                out.image_count += len(page.images or [])
                for ch in page.chars or []:
                    if fn := ch.get("fontname"):
                        out.fonts.add(str(fn).split("+")[-1])
                if _detect_pdf_columns(words, float(page.width), float(page.height)):
                    out.has_multi_column = True
                out.header_footer_text.extend(
                    _pdf_edge_text(words, float(page.height))
                )
            except Exception as exc:  # noqa: BLE001
                out.parse_errors.append(f"page parse failed: {exc}")

    out.text = "\n".join(chunks).strip()
    # A PDF with pages but effectively no extractable text is a scanned image.
    out.has_text_layer = len(out.text) >= 50
    return out


def _detect_pdf_columns(words: list[dict], width: float, height: float) -> bool:
    """Detect a genuine two-column layout via a persistent vertical gutter.

    Looks for a band of x-positions in the middle of the page that no word
    crosses, while substantial text sits on both sides. Checked against word
    spans rather than a naive left/right word count, which would flag ordinary
    right-aligned dates as a column.
    """
    if len(words) < 40:
        return False
    bins = 50
    occupied = [False] * bins
    for w in words:
        try:
            x0, x1 = float(w["x0"]), float(w["x1"])
        except (KeyError, TypeError, ValueError):
            continue
        lo = max(0, min(bins - 1, int(x0 / width * bins)))
        hi = max(0, min(bins - 1, int(x1 / width * bins)))
        for i in range(lo, hi + 1):
            occupied[i] = True

    # Only the middle 30-70% of the width can be a real gutter.
    lo_b, hi_b = int(bins * 0.30), int(bins * 0.70)
    run = best_run = 0
    for i in range(lo_b, hi_b):
        run = 0 if occupied[i] else run + 1
        best_run = max(best_run, run)
    if best_run < 3:
        return False

    left = sum(1 for w in words if float(w.get("x1", 0)) < width * 0.45)
    right = sum(1 for w in words if float(w.get("x0", width)) > width * 0.55)
    # Both sides must carry real content, not just a sidebar of icons.
    return left >= 15 and right >= 15


def _pdf_edge_text(words: list[dict], height: float) -> list[str]:
    """Text in the top/bottom 6% of the page, which many ATS parsers discard."""
    top_band, bottom_band = height * 0.06, height * 0.94
    found: list[str] = []
    for w in words:
        try:
            top, bottom = float(w["top"]), float(w["bottom"])
        except (KeyError, TypeError, ValueError):
            continue
        if top < top_band or bottom > bottom_band:
            if text := str(w.get("text", "")).strip():
                found.append(text)
    return found


# ---------------------------------------------------------------- docx


def _parse_docx(p: Path) -> ParsedResume:
    import docx

    out = ParsedResume(path=str(p), file_type="docx")
    doc = docx.Document(str(p))

    lines = [para.text for para in doc.paragraphs]

    for table in doc.tables:
        out.table_count += 1
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)

    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is None:
                continue
            for para in part.paragraphs:
                if para.text.strip():
                    out.header_footer_text.append(para.text.strip())

    try:
        out.image_count = len(doc.inline_shapes)
    except Exception:  # noqa: BLE001
        pass

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font is not None and run.font.name:
                out.fonts.add(run.font.name)

    # A docx has no page count without rendering; approximate for length checks.
    out.text = "\n".join(lines).strip()
    out.pages = max(1, round(out.word_count / 500)) if out.word_count else 0
    out.has_text_layer = bool(out.text)
    # Multi-column in docx is usually implemented as a borderless layout table.
    out.has_multi_column = _docx_layout_table(doc)
    return out


def _docx_layout_table(doc) -> bool:  # noqa: ANN001
    """A wide 2-3 column table holding most of the document is a layout hack."""
    for table in doc.tables:
        cols = len(table.columns) if table.columns else 0
        if 2 <= cols <= 3:
            chars = sum(len(c.text) for r in table.rows for c in r.cells)
            if chars > 400:
                return True
    return False


# ---------------------------------------------------------------- plain text


def _parse_text(p: Path) -> ParsedResume:
    text = p.read_text(encoding="utf-8", errors="replace").strip()
    return ParsedResume(
        path=str(p),
        file_type=p.suffix.lower().lstrip("."),
        text=text,
        pages=max(1, round(len(text.split()) / 500)) if text else 0,
        has_text_layer=bool(text),
    )


# ---------------------------------------------------------------- content


def _extract_content_signals(out: ParsedResume) -> None:
    out.emails = list(dict.fromkeys(_EMAIL.findall(out.text)))
    out.urls = list(dict.fromkeys(_URL.findall(out.text)))
    out.phones = [
        m.group(0).strip()
        for m in _PHONE.finditer(out.text)
        if sum(c.isdigit() for c in m.group(0)) >= 9
    ]

    alias_to_key = {
        alias: key for key, aliases in CANONICAL_SECTIONS.items() for alias in aliases
    }

    for raw_line in out.text.splitlines():
        line = raw_line.strip()
        if not line or len(line) > 60:
            continue
        probe = re.sub(r"[^a-z\s&/]", "", line.lower()).strip()
        if not probe:
            continue
        if key := alias_to_key.get(probe):
            out.sections_found.setdefault(key, line)
            continue
        # Heading-shaped (short, title/upper case, no sentence punctuation) but
        # not recognisable: a likely ATS misfile.
        words = line.split()
        looks_heading = (
            1 <= len(words) <= 5
            and not line.endswith((".", ",", ";", ":"))
            and (line.isupper() or line.istitle())
            and not _EMAIL.search(line)
            and not any(ch.isdigit() for ch in line)
        )
        if looks_heading and probe not in alias_to_key:
            out.unknown_headings.append(line)

    for glyph in ("•", "▪", "◦", "‣", "·", "»", "→", "✦", "★", "-", "*"):
        if glyph in out.text:
            out.bullet_glyphs.add(glyph)
