"""Render tailored resume text to .docx and .pdf.

Deliberately plain: single column, no tables, no text boxes, no headers or
footers, one safe font. Every one of those is a thing the rule engine in
`app/pipeline/ats_rules.py` flags, so generating them here would be
self-defeating. Visual polish is intentionally sacrificed for parseability.

.docx is the preferred upload format for most ATS platforms; PDF is produced as a
human-readable copy.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from app.config import get_settings

log = logging.getLogger(__name__)

FONT = "Calibri"
FONT_PT = 11
NAME_PT = 16
HEADING_PT = 12

_SECTION_WORDS = {
    "summary", "professional summary", "profile", "objective", "about",
    "experience", "work experience", "professional experience", "employment",
    "employment history", "work history", "education", "skills",
    "technical skills", "core competencies", "projects", "certifications",
    "publications", "awards", "languages", "interests", "volunteering",
}

_BULLET = re.compile(r"^\s*[-*•‣▪●·]\s+")


def slugify(value: str, *, limit: int = 40) -> str:
    out = re.sub(r"[^a-zA-Z0-9]+", "_", value or "").strip("_").lower()
    return (out or "job")[:limit]


def _is_heading(line: str) -> bool:
    stripped = line.strip().rstrip(":")
    if not stripped or len(stripped) > 40 or _BULLET.match(line):
        return False
    if stripped.lower() in _SECTION_WORDS:
        return True
    # ALL CAPS short lines are conventionally headings.
    return stripped.isupper() and len(stripped.split()) <= 4


def write_docx(text: str, path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(FONT_PT)
    # Word needs east-asian set explicitly or it silently substitutes a font.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT
    )

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(50)

    lines = (text or "").splitlines()
    first_written = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue

        if not first_written:
            para = doc.add_paragraph()
            run = para.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(NAME_PT)
            first_written = True
            continue

        if _is_heading(line):
            para = doc.add_paragraph()
            run = para.add_run(line.strip().rstrip(":").upper())
            run.bold = True
            run.font.size = Pt(HEADING_PT)
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(2)
            continue

        if _BULLET.match(line):
            # Use a real list style so the bullet is structural, not a literal glyph.
            para = doc.add_paragraph(_BULLET.sub("", line).strip(),
                                     style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            continue

        para = doc.add_paragraph(line.strip())
        para.paragraph_format.space_after = Pt(2)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def write_pdf(text: str, path: Path) -> Path:
    from fpdf import FPDF

    pdf = FPDF(unit="pt", format="A4")
    pdf.set_auto_page_break(auto=True, margin=40)
    pdf.add_page()
    pdf.set_margins(50, 40, 50)

    usable = pdf.w - 100
    first_written = False

    for raw in (text or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            pdf.ln(6)
            continue

        # Core fonts are latin-1 only; drop what cannot be encoded rather than crash.
        safe = line.encode("latin-1", "replace").decode("latin-1")

        if not first_written:
            pdf.set_font("Helvetica", "B", NAME_PT)
            pdf.multi_cell(usable, NAME_PT + 5, safe.strip())
            first_written = True
            continue

        if _is_heading(line):
            pdf.ln(8)
            pdf.set_font("Helvetica", "B", HEADING_PT)
            pdf.multi_cell(usable, HEADING_PT + 3, safe.strip().rstrip(":").upper())
            continue

        pdf.set_font("Helvetica", "", FONT_PT)
        if _BULLET.match(line):
            body = _BULLET.sub("", safe).strip()
            pdf.multi_cell(usable, FONT_PT + 4, f"•  {body}"
                           .encode("latin-1", "replace").decode("latin-1"))
        else:
            pdf.multi_cell(usable, FONT_PT + 4, safe.strip())

    path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(path))
    return path


def render_for_job(
    text: str, *, company: str, title: str, full_name: str = ""
) -> tuple[Path, Path]:
    """Write both formats into the artifact dir. Returns (docx, pdf).

    A PDF failure is tolerated: the .docx is what gets submitted, so losing the
    human-readable copy should not fail an application.
    """
    settings = get_settings()
    settings.ensure_dirs()

    stem = "_".join(
        p for p in (slugify(full_name or "resume", limit=24),
                    slugify(company, limit=24),
                    slugify(title, limit=32)) if p
    )
    base = settings.artifact_dir / stem

    docx_path = write_docx(text, base.with_suffix(".docx"))
    try:
        pdf_path = write_pdf(text, base.with_suffix(".pdf"))
    except Exception as exc:  # noqa: BLE001
        log.warning("PDF render failed for %s (docx is unaffected): %s", stem, exc)
        pdf_path = Path()

    return docx_path, pdf_path
